"""Extract text + light metadata from customer documents.

Supported: .pdf (pdfplumber -> pypdf fallback), .docx (python-docx),
.xlsx/.xls (openpyxl / pandas+xlrd), .csv (pandas).  Everything is best-effort:
a failure on one file is captured in ``ExtractedDoc.error`` rather than raised.
"""
from __future__ import annotations

import re
import uuid
from pathlib import Path

from ..schemas import ExtractedDoc, TablePreview

_MAX_TEXT_CHARS = 20_000
_EXCERPT_CHARS = 1_600
_MAX_PDF_PAGES = 25
_MAX_SHEET_ROWS = 40
_MAX_TABLE_ROWS = 8

_STOPWORDS = set(
    """the a an and or of for to in on at by with from as is are be this that these those
    all any per your our their its it we you they he she not no if then than into out over
    under more most such other same each which who whom whose will shall may can copy
    original certified signed current valid dated date list schedule report statement
    document documents evidence type provided pending remarks section
    xlsx xls csv pdf docx doc sheet sheet1 worksheet workbook unnamed nan none
    name society residence""".split()
)


def _tokens(text: str) -> list[str]:
    words = re.findall(r"[a-z][a-z0-9&/\-]{2,}", text.lower())
    return [w.strip("-/&") for w in words if w not in _STOPWORDS and len(w) > 2]


def _keywords(text: str, filename: str, folder: str, k: int = 25) -> list[str]:
    from collections import Counter

    bag = _tokens(f"{filename} {folder} {folder} {text}")
    common = [w for w, _ in Counter(bag).most_common(k * 2)]
    seen, out = set(), []
    for w in common:
        if w not in seen:
            seen.add(w)
            out.append(w)
        if len(out) >= k:
            break
    return out


# --------------------------------------------------------------------------- #
# per-format readers
# --------------------------------------------------------------------------- #
def _read_pdf(path: Path) -> tuple[str, int]:
    text_parts: list[str] = []
    pages = 0
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages[:_MAX_PDF_PAGES]:
                text_parts.append(page.extract_text() or "")
        joined = "\n".join(text_parts).strip()
        if joined:
            return joined[:_MAX_TEXT_CHARS], pages
    except Exception:
        pass
    # fallback
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = len(reader.pages)
        for page in reader.pages[:_MAX_PDF_PAGES]:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts).strip()[:_MAX_TEXT_CHARS], pages
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"PDF read failed: {exc}") from exc


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:_MAX_TEXT_CHARS]


def _read_xlsx(path: Path) -> tuple[str, list[str], list[TablePreview]]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    sheet_names = list(wb.sheetnames)
    parts: list[str] = []
    tables: list[TablePreview] = []
    for sn in sheet_names:
        ws = wb[sn]
        rows_txt: list[list[str]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= _MAX_SHEET_ROWS:
                break
            vals = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in vals):
                rows_txt.append(vals)
        if not rows_txt:
            continue
        parts.append(f"# Sheet: {sn}")
        for r in rows_txt:
            parts.append(" | ".join(v for v in r if v.strip()))
        width = max(len(r) for r in rows_txt)
        tables.append(
            TablePreview(
                sheet=sn,
                columns=[f"C{c+1}" for c in range(min(width, 12))],
                rows=[r[:12] for r in rows_txt[:_MAX_TABLE_ROWS]],
            )
        )
    wb.close()
    return "\n".join(parts)[:_MAX_TEXT_CHARS], sheet_names, tables


def _read_xls_csv(path: Path) -> tuple[str, list[str], list[TablePreview]]:
    import pandas as pd

    parts: list[str] = []
    tables: list[TablePreview] = []
    sheets: list[str] = []
    if path.suffix.lower() == ".csv":
        frames = {"CSV": pd.read_csv(path, dtype=str, nrows=500, on_bad_lines="skip")}
    else:  # .xls
        frames = pd.read_excel(path, sheet_name=None, dtype=str, nrows=200, engine="xlrd")
    for sn, df in frames.items():
        sheets.append(sn)
        df = df.fillna("")
        parts.append(f"# Sheet: {sn}")
        parts.append(" | ".join(str(c) for c in df.columns))
        for _, row in df.head(_MAX_SHEET_ROWS).iterrows():
            vals = [str(v) for v in row.tolist() if str(v).strip()]
            if vals:
                parts.append(" | ".join(vals))
        tables.append(
            TablePreview(
                sheet=sn,
                columns=[str(c) for c in df.columns[:12]],
                rows=[[str(v) for v in r[:12]] for r in df.head(_MAX_TABLE_ROWS).values.tolist()],
            )
        )
    return "\n".join(parts)[:_MAX_TEXT_CHARS], sheets, tables


# --------------------------------------------------------------------------- #
# public entry point
# --------------------------------------------------------------------------- #
def extract_document(path: Path, rel_path: str) -> ExtractedDoc:
    rel_path = rel_path.replace("\\", "/").lstrip("/")
    parts = [p for p in rel_path.split("/") if p]
    # folder = the segment that names a section if present, else the file's
    # immediate parent directory (so selecting a wrapper folder still keeps the
    # meaningful "SECTION n" / "5.2" grouping).
    section_seg = next((p for p in parts[:-1] if re.search(r"section\s*[-\s]*\d", p, re.I)), None)
    folder = section_seg or (parts[-2] if len(parts) >= 2 else "")
    ext = path.suffix.lower()
    doc = ExtractedDoc(
        id=uuid.uuid4().hex[:12],
        filename=path.name,
        rel_path=rel_path,
        folder=folder,
        ext=ext,
        size_bytes=path.stat().st_size if path.exists() else 0,
    )
    try:
        text = ""
        if ext == ".pdf":
            text, doc.page_count = _read_pdf(path)
        elif ext == ".docx":
            text = _read_docx(path)
        elif ext == ".doc":
            doc.error = ".doc (legacy Word) not supported - please resave as .docx"
        elif ext == ".xlsx":
            text, doc.sheet_names, doc.tables = _read_xlsx(path)
        elif ext in (".xls", ".csv"):
            text, doc.sheet_names, doc.tables = _read_xls_csv(path)
        else:
            doc.error = f"unsupported extension {ext}"

        text = text or ""
        doc.char_count = len(text)
        doc.text_excerpt = text[:_EXCERPT_CHARS]
        doc.keywords = _keywords(text, path.name, folder)
        if not text and not doc.error:
            doc.error = "no extractable text (scanned image or empty file)"
    except Exception as exc:  # noqa: BLE001 - surface, don't crash the batch
        doc.error = f"{type(exc).__name__}: {exc}"
    return doc


def full_text(path: Path) -> str:
    """Re-read a document's full (capped) text - used by the LLM pass."""
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _read_pdf(path)[0]
        if ext == ".docx":
            return _read_docx(path)
        if ext == ".xlsx":
            return _read_xlsx(path)[0]
        if ext in (".xls", ".csv"):
            return _read_xls_csv(path)[0]
    except Exception:
        return ""
    return ""


def full_table_text(path: Path, max_rows: int = 220, max_cols: int = 12) -> str:
    """Dump a spreadsheet's rows as pipe-delimited text for LLM extraction -
    far more rows than the on-screen excerpt (a Mollak trial balance is ~450)."""
    ext = path.suffix.lower()
    parts: list[str] = []
    try:
        if ext == ".xlsx":
            import openpyxl

            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
            for sn in wb.sheetnames:
                ws = wb[sn]
                parts.append(f"# sheet: {sn}")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= max_rows:
                        parts.append("# ...(truncated)")
                        break
                    cells = [
                        "" if v is None else (f"{v:.2f}" if isinstance(v, float) else str(v))
                        for v in row[:max_cols]
                    ]
                    if any(c.strip() for c in cells):
                        parts.append(" | ".join(cells).rstrip(" |"))
            wb.close()
        elif ext in (".xls", ".csv"):
            import pandas as pd

            frames = (
                {"csv": pd.read_csv(path, dtype=str, header=None, nrows=max_rows, on_bad_lines="skip")}
                if ext == ".csv"
                else pd.read_excel(path, sheet_name=None, dtype=str, header=None, engine="xlrd")
            )
            for sn, df in frames.items():
                parts.append(f"# sheet: {sn}")
                for row in df.head(max_rows).values.tolist():
                    cells = ["" if (v is None or str(v) == "nan") else str(v) for v in row[:max_cols]]
                    if any(c.strip() for c in cells):
                        parts.append(" | ".join(cells).rstrip(" |"))
        else:
            return full_text(path)[:12000]
    except Exception as exc:  # noqa: BLE001
        return f"(could not read: {exc})"
    return "\n".join(parts)[:16000]
